import streamlit as st
from transformers import MarianMTModel, MarianTokenizer
import docx
import pdfplumber
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import tempfile
import os

os.environ["TF_CPP_MIN_LOG_LEVEL"] = "2"  # masque messages TensorFlow
os.environ["TRANSFORMERS_CACHE"] = "./cache"  # cache HuggingFace

st.set_page_config(page_title="---------------- Traduction de documents-------------", layout="wide")

def read_word(file_path):
    doc = docx.Document(file_path)
    return [p.text for p in doc.paragraphs if p.text.strip() != ""]

def read_pdf(file_path):
    text_list = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text_list += [line.strip() for line in page.extract_text().split("\n") if line.strip() != ""]
    return text_list


def translate_text(paragraphs, src_lang="fr", tgt_lang="en", batch_size=5):
    model_name = f"Helsinki-NLP/opus-mt-{src_lang}-{tgt_lang}"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    
    translated_paragraphs = []
    progress_bar = st.progress(0)
    total = len(paragraphs)
    
    for i in range(0, total, batch_size):
        batch = tokenizer(paragraphs[i:i+batch_size], return_tensors="pt", padding=True)
        gen = model.generate(**batch)
        translated_batch = [tokenizer.decode(g, skip_special_tokens=True) for g in gen]
        translated_paragraphs.extend(translated_batch)
        progress_bar.progress(min((i+batch_size)/total, 1.0))
    
    progress_bar.empty()
    return translated_paragraphs

def save_to_word(paragraphs, output_path):
    doc = docx.Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output_path)

def save_to_pdf(paragraphs, output_path):
    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    story = []

    for para in paragraphs:
        story.append(Paragraph(para, styles["Normal"]))
        story.append(Spacer(1, 5))  # espace entre paragraphes
    
    doc.build(story)

st.title("******* Traduction automatique de documents")
st.write("Upload un document Word ou PDF et choisissez la langue de traduction.")

with st.expander(" Traduction de documents"):
    uploaded_file = st.file_uploader("Choisir un fichier", type=["pdf", "docx"])
    src_lang = st.selectbox("Langue source", ["fr", "en", "es", "de", "it"])
    tgt_lang = st.selectbox("Langue cible", ["fr", "en", "es", "de", "it"])
    
    if uploaded_file and st.button("Traduire et Générer Word + PDF"):
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name
        
        if uploaded_file.type == "application/pdf":
            paragraphs = read_pdf(tmp_path)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            paragraphs = read_word(tmp_path)
        else:
            st.error("Format non supporté !")
            st.stop()
        
        st.info("Traduction en cours...")
        translated_paragraphs = translate_text(paragraphs, src_lang=src_lang, tgt_lang=tgt_lang)
        
        # Sauvegarde Word + PDF
        word_file = "document_traduit.docx"
        pdf_file = "document_traduit.pdf"
        save_to_word(translated_paragraphs, word_file)
        save_to_pdf(translated_paragraphs, pdf_file)
        
        st.success("Traduction terminée !")
        st.download_button("Télécharger Word", data=open(word_file, "rb").read(), file_name=word_file)
        st.download_button(" Télécharger PDF", data=open(pdf_file, "rb").read(), file_name=pdf_file)
