import streamlit as st
from transformers import MarianMTModel, MarianTokenizer
import docx
import pdfplumber
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import tempfile
import os

# =======================
# Configuration générale
# =======================
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "2"  # Masquer messages TensorFlow
os.environ["TRANSFORMERS_CACHE"] = "./cache"  # Cache HuggingFace

st.set_page_config(page_title="Traduction de documents académique", layout="wide")

# =======================
# Fonctions de lecture
# =======================
def read_word(file_path: str) -> list[str]:
    doc = docx.Document(file_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def read_pdf(file_path: str) -> list[str]:
    paragraphs = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                for line in page.extract_text().split("\n"):
                    line = line.strip()
                    if line:
                        paragraphs.append(line)
    return paragraphs

# =======================
# Détection titre
# =======================
def is_title(paragraph: str) -> bool:
    return paragraph.isupper() and len(paragraph.split()) < 10

# =======================
# Traduction par paragraphe
# =======================
def translate_paragraphs(paragraphs: list[str], src_lang="fr", tgt_lang="en") -> list[str]:
    model_name = f"Helsinki-NLP/opus-mt-{src_lang}-{tgt_lang}"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)

    translated = []
    progress_bar = st.progress(0)
    total = len(paragraphs)

    for i, para in enumerate(paragraphs):
        # Eviter de tronquer le texte
        batch = tokenizer([para], return_tensors="pt", padding=True, truncation=True, max_length=1024)
        generated = model.generate(**batch)
        translated_text = tokenizer.decode(generated[0], skip_special_tokens=True)
        translated.append(translated_text)
        progress_bar.progress((i+1)/total)

    progress_bar.empty()
    return translated

# =======================
# Sauvegarde Word / PDF
# =======================
def save_as_word(paragraphs: list[str], output_path: str):
    doc = docx.Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output_path)

def save_as_pdf(paragraphs: list[str], output_path: str):
    doc = SimpleDocTemplate(output_path)
    styles = getSampleStyleSheet()
    story = []

    for para in paragraphs:
        story.append(Paragraph(para, styles["Normal"]))
        story.append(Spacer(1, 5))

    doc.build(story)

# =======================
# Interface Streamlit
# =======================
st.title("Traduction professionnelle de documents")
st.write("Téléversez un document Word ou PDF et choisissez la langue source et cible.")

uploaded_file = st.file_uploader("Sélectionner un fichier", type=["pdf", "docx"])
src_lang = st.selectbox("Langue source", ["fr", "en", "es", "de", "it"])
tgt_lang = st.selectbox("Langue cible", ["fr", "en", "es", "de", "it"])

if uploaded_file and st.button("Traduire et générer Word + PDF"):
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_path = tmp_file.name

    # Lecture
    if uploaded_file.type == "application/pdf":
        paragraphs = read_pdf(tmp_path)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        paragraphs = read_word(tmp_path)
    else:
        st.error("Format non supporté !")
        st.stop()

    st.info("Traduction en cours…")
    translated_paragraphs = translate_paragraphs(paragraphs, src_lang, tgt_lang)

    # Sauvegarde
    word_file = "document_traduit.docx"
    pdf_file = "document_traduit.pdf"
    save_as_word(translated_paragraphs, word_file)
    save_as_pdf(translated_paragraphs, pdf_file)

    st.success("Traduction terminée !")
    st.download_button("Télécharger Word", data=open(word_file, "rb").read(), file_name=word_file)
    st.download_button("Télécharger PDF", data=open(pdf_file, "rb").read(), file_name=pdf_file)
