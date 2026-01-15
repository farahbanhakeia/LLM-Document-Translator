import streamlit as st
from transformers import MarianMTModel, MarianTokenizer
import docx
import pdfplumber
import tempfile

def read_word(file_path):
    doc = docx.Document(file_path)
    return [p.text for p in doc.paragraphs if p.text.strip() != ""]

def read_pdf(file_path):
    text_list = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text_list += [line.strip() for line in page.extract_text().split("\n") if line.strip() != ""]
    return text_list

def translate_text(paragraphs, src_lang="fr", tgt_lang="en"):
    model_name = f"Helsinki-NLP/opus-mt-{src_lang}-{tgt_lang}"
    tokenizer = MarianTokenizer.from_pretrained(model_name)
    model = MarianMTModel.from_pretrained(model_name)
    
    translated_paragraphs = []
    for paragraph in paragraphs:
        batch = tokenizer([paragraph], return_tensors="pt", padding=True)
        gen = model.generate(**batch)
        translated = tokenizer.decode(gen[0], skip_special_tokens=True)
        translated_paragraphs.append(translated)
    
    return translated_paragraphs

def save_to_word(paragraphs, output_path):
    doc = docx.Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output_path)

st.title("----------Traduction automatique de documents---------------------")
st.write("Upload un document Word ou PDF et choisissez la langue de traduction.")

uploaded_file = st.file_uploader("Choisir un fichier", type=["pdf", "docx"])
src_lang = st.selectbox("Langue source", ["fr", "en", "es", "de", "it"])
tgt_lang = st.selectbox("Langue cible", ["fr", "en", "es", "de", "it"])

if uploaded_file and st.button("Traduire et Générer Word"):
    # Sauvegarder le fichier temporairement
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_path = tmp_file.name
    
    # Lire le document
    if uploaded_file.type == "application/pdf":
        paragraphs = read_pdf(tmp_path)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        paragraphs = read_word(tmp_path)
    else:
        st.error("Format non supporté !")
        st.stop()
    
    st.info("Traduction en cours... ça peut prendre quelques secondes selon la taille du document.")
    translated_paragraphs = translate_text(paragraphs, src_lang=src_lang, tgt_lang=tgt_lang)
    
    # Sauvegarder le Word traduit
    output_file = "document_traduit.docx"
    save_to_word(translated_paragraphs, output_file)
    
    st.success(f"Traduction terminée ! Fichier généré : {output_file}")
    st.download_button(" Télécharger le document traduit", data=open(output_file, "rb").read(), file_name=output_file)
